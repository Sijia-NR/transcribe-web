#!/usr/bin/env python3
"""讯飞录音转写 Web 应用"""

import hashlib
import hmac
import base64
import urllib.parse
import random
import string
import time
import json
import os
import uuid
import threading
from datetime import datetime, timezone, timedelta

import requests
from flask import (
    Flask, render_template, request, redirect, url_for,
    session, send_file, jsonify
)
from docx import Document
from docx.shared import Pt

# === 配置 ===
BASE_URL = "https://office-api-ist-dx.iflyaisol.com"
SECRET_KEY = os.environ.get("SECRET_KEY", "transcribe-web-secret-change-me")
APP_PASSWORD = os.environ.get("APP_PASSWORD", "1111iran")
PORT = int(os.environ.get("PORT", 9023))

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
RESULT_DIR = os.path.join(BASE_DIR, "results")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(RESULT_DIR, exist_ok=True)

app = Flask(__name__)
app.secret_key = SECRET_KEY

# 转写任务状态（内存存储）
tasks = {}
tasks_lock = threading.Lock()


# ========== 讯飞 API 核心逻辑 ==========

def generate_signature(params, secret):
    sorted_keys = sorted(k for k in params.keys() if k != "signature")
    pairs = []
    for k in sorted_keys:
        v = params[k]
        if v is None or v == "":
            continue
        encoded_key = urllib.parse.quote(str(k), safe="")
        encoded_value = urllib.parse.quote(str(v), safe="")
        pairs.append(f"{encoded_key}={encoded_value}")
    base_string = "&".join(pairs)
    digest = hmac.new(
        secret.encode("utf-8"),
        base_string.encode("utf-8"),
        hashlib.sha1,
    ).digest()
    return base64.b64encode(digest).decode("utf-8")


def generate_random_string(length=16):
    chars = string.ascii_letters + string.digits
    return "".join(random.choice(chars) for _ in range(length))


def build_url(path, params):
    query = urllib.parse.urlencode(params)
    return f"{BASE_URL}{path}?{query}"


def get_audio_duration_ms(filepath):
    import subprocess
    result = subprocess.run(
        [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            filepath,
        ],
        capture_output=True, text=True,
    )
    try:
        seconds = float(result.stdout.strip())
    except (ValueError, TypeError):
        seconds = 0
    return int(seconds * 1000)


def upload_audio(app_id, api_key, api_secret, filepath):
    """上传音频文件到讯飞"""
    file_size = os.path.getsize(filepath)
    duration_ms = get_audio_duration_ms(filepath)
    file_name = os.path.basename(filepath)

    tz = timezone(timedelta(hours=8))
    date_time = datetime.now(tz).strftime("%Y-%m-%dT%H:%M:%S+0800")
    sig_random = generate_random_string()

    params = {
        "appId": app_id,
        "accessKeyId": api_key,
        "dateTime": date_time,
        "signatureRandom": sig_random,
        "fileSize": str(file_size),
        "fileName": file_name,
        "duration": str(duration_ms),
        "language": "autodialect",
        "roleType": "1",
        "roleNum": "0",
    }

    signature = generate_signature(params, api_secret)
    url = build_url("/v2/upload", params)
    headers = {
        "Content-Type": "application/octet-stream",
        "signature": signature,
    }

    with open(filepath, "rb") as f:
        resp = requests.post(url, headers=headers, data=f)

    result = resp.json()
    print(f"[DEBUG] upload response: {json.dumps(result, ensure_ascii=False)}")
    if result.get("code") != "000000":
        raise Exception(f"上传失败: {result.get('descInfo', '未知错误')}")

    order_id = result["content"]["orderId"]
    estimate_time = result["content"].get("taskEstimateTime", 0)
    return order_id, estimate_time


def poll_result(api_key, api_secret, order_id):
    """轮询转写结果（单次查询）"""
    tz = timezone(timedelta(hours=8))
    date_time = datetime.now(tz).strftime("%Y-%m-%dT%H:%M:%S+0800")
    sig_random = generate_random_string()

    params = {
        "accessKeyId": api_key,
        "dateTime": date_time,
        "signatureRandom": sig_random,
        "orderId": order_id,
        "resultType": "transfer",
    }

    signature = generate_signature(params, api_secret)
    url = build_url("/v2/getResult", params)
    headers = {
        "Content-Type": "application/json",
        "signature": signature,
    }

    resp = requests.post(url, headers=headers, json={}, timeout=30)
    result = resp.json()

    content = result.get("content", {})
    order_info = content.get("orderInfo", {})
    status = order_info.get("status", -1)

    # 调试：状态为完成时打印原始结果
    if status == 4:
        print(f"[DEBUG] getResult orderInfo: {json.dumps(order_info, ensure_ascii=False)}")
        raw = content.get("orderResult", "")
        if raw:
            print(f"[DEBUG] orderResult (first 2000 chars): {str(raw)[:2000]}")

    return status, content


def parse_result(content):
    """解析转写结果"""
    order_result_str = content.get("orderResult")
    if not order_result_str:
        return []

    order_result = (
        json.loads(order_result_str)
        if isinstance(order_result_str, str)
        else order_result_str
    )
    lattice_list = order_result.get("lattice", [])

    segments = []
    for lattice in lattice_list:
        json_1best = lattice.get("json_1best", {})
        if isinstance(json_1best, str):
            json_1best = json.loads(json_1best)

        st = json_1best.get("st", {})
        rl = st.get("rl", "未知")

        rt_list = st.get("rt", [])
        ws_list = rt_list[0].get("ws", []) if rt_list else st.get("ws", [])

        words = []
        for ws in ws_list:
            for cw in ws.get("cw", []):
                wp = cw.get("wp", "")
                if wp != "s":
                    words.append(cw.get("w", ""))

        if words:
            segments.append((rl, "".join(words)))

    # 合并同一说话人连续段落
    merged = []
    for rl, text in segments:
        if merged and merged[-1][0] == rl:
            merged[-1] = (rl, merged[-1][1] + text)
        else:
            merged.append((rl, text))

    return merged


# ========== 讯飞极速版 API ==========

OST_UPLOAD_URL = "https://upload-ost-api.xfyun.cn/file/upload"
OST_MPINIT_URL = "https://upload-ost-api.xfyun.cn/file/mpupload/init"
OST_MPUPLOAD_URL = "https://upload-ost-api.xfyun.cn/file/mpupload/upload"
OST_MPCOMPLETE_URL = "https://upload-ost-api.xfyun.cn/file/mpupload/complete"
OST_CREATE_URL = "https://ost-api.xfyun.cn/v2/ost/pro_create"
OST_QUERY_URL = "https://ost-api.xfyun.cn/v2/ost/query"


def ost_build_auth(url, method, api_key, api_secret, body_bytes, use_empty_digest=False):
    """构建极速版 HMAC-SHA256 鉴权

    注意：文件上传接口的 digest 需要使用空字符串计算（与讯飞官方 demo 一致），
    创建/查询任务接口则使用实际 body 计算 digest。
    """
    from email.utils import formatdate

    parsed = urllib.parse.urlparse(url)
    host = parsed.hostname
    date = formatdate(usegmt=True)
    request_line = f"{method.upper()} {parsed.path} HTTP/1.1"

    if use_empty_digest:
        # 文件上传接口：digest 使用空字符串（官方 demo 行为）
        body_digest = hashlib.sha256(b"").digest()
    else:
        body_digest = hashlib.sha256(body_bytes).digest()
    digest = "SHA-256=" + base64.b64encode(body_digest).decode("utf-8")

    signature_origin = (
        f"host: {host}\ndate: {date}\n"
        f"{request_line}\ndigest: {digest}"
    )
    signature_hmac = hmac.new(
        api_secret.encode("utf-8"),
        signature_origin.encode("utf-8"),
        hashlib.sha256,
    ).digest()
    signature = base64.b64encode(signature_hmac).decode("utf-8")

    authorization = (
        f'api_key="{api_key}", algorithm="hmac-sha256", '
        f'headers="host date request-line digest", signature="{signature}"'
    )
    return {
        "host": host,
        "date": date,
        "digest": digest,
        "authorization": authorization,
    }


def _ost_multipart_upload(url, app_id, api_key, api_secret, filepath, extra_fields=None):
    """通用 multipart 上传，返回响应 JSON"""
    request_id = str(uuid.uuid4())
    file_name = os.path.basename(filepath)

    boundary = uuid.uuid4().hex

    parts = [
        (f'--{boundary}\r\nContent-Disposition: form-data; name="request_id"'
         f"\r\n\r\n{request_id}\r\n").encode(),
        (f'--{boundary}\r\nContent-Disposition: form-data; name="app_id"'
         f"\r\n\r\n{app_id}\r\n").encode(),
    ]
    if extra_fields:
        for k, v in extra_fields.items():
            parts.append(
                (f'--{boundary}\r\nContent-Disposition: form-data; name="{k}"'
                 f"\r\n\r\n{v}\r\n").encode()
            )
    with open(filepath, "rb") as f:
        file_content = f.read()
    parts.append(
        (f'--{boundary}\r\nContent-Disposition: form-data; name="data"; '
         f'filename="{file_name}"\r\nContent-Type: application/octet-stream\r\n\r\n').encode()
    )
    body = b"".join(parts) + file_content + f"\r\n--{boundary}--\r\n".encode()

    headers = ost_build_auth(url, "POST", api_key, api_secret, b"", use_empty_digest=True)
    headers["content-type"] = f"multipart/form-data; boundary={boundary}"

    resp = requests.post(url, headers=headers, data=body, timeout=(30, 600))
    return resp.json()


def ost_upload_file(app_id, api_key, api_secret, filepath):
    """极速版：上传音频文件，返回 (audio_url, 格式)"""
    ext = os.path.splitext(filepath)[1].lower().lstrip(".")
    if ext not in ("wav", "mp3", "pcm"):
        import subprocess

        wav_path = filepath + ".wav"
        subprocess.run(
            [
                "ffmpeg", "-y", "-i", filepath,
                "-ar", "16000", "-ac", "1", "-sample_fmt", "s16", wav_path,
            ],
            capture_output=True,
        )
        if os.path.exists(wav_path):
            filepath = wav_path
            ext = "wav"

    file_size = os.path.getsize(filepath)

    if file_size > 30 * 1024 * 1024:
        # 大文件走分块上传
        audio_url = ost_upload_chunks(app_id, api_key, api_secret, filepath)
        return audio_url, ext

    # 小文件直接上传
    result = _ost_multipart_upload(OST_UPLOAD_URL, app_id, api_key, api_secret, filepath)
    print(f"[DEBUG] OST upload response: {json.dumps(result, ensure_ascii=False)}")

    if result.get("code") != 0:
        raise Exception(f"上传失败: {result.get('message', '未知错误')}")

    return result["data"]["url"], ext


def ost_upload_chunks(app_id, api_key, api_secret, filepath):
    """极速版：大文件分块上传（>30M）"""
    # 1) 初始化分块
    body = {"request_id": str(uuid.uuid4()), "app_id": app_id}
    body_bytes = json.dumps(body).encode("utf-8")
    headers = ost_build_auth(OST_MPINIT_URL, "POST", api_key, api_secret, body_bytes)
    headers["content-type"] = "application/json"
    resp = requests.post(OST_MPINIT_URL, headers=headers, data=body_bytes, timeout=30)
    result = resp.json()
    print(f"[DEBUG] OST mpinit response: {json.dumps(result, ensure_ascii=False)}")
    if result.get("code") != 0:
        raise Exception(f"分块初始化失败: {result.get('message', '未知错误')}")
    upload_id = result["data"]["upload_id"]

    # 2) 分块上传（每块 10MB）
    chunk_size = 10 * 1024 * 1024
    slice_id = 0
    with open(filepath, "rb") as f:
        while True:
            chunk_data = f.read(chunk_size)
            if not chunk_data:
                break
            # 写临时文件用于 multipart 上传
            tmp_path = filepath + f".part{slice_id}"
            with open(tmp_path, "wb") as tf:
                tf.write(chunk_data)
            try:
                r = _ost_multipart_upload(
                    OST_MPUPLOAD_URL, app_id, api_key, api_secret, tmp_path,
                    extra_fields={"upload_id": upload_id, "slice_id": str(slice_id)},
                )
                print(f"[DEBUG] OST chunk {slice_id} response: {json.dumps(r, ensure_ascii=False)}")
                if r.get("code") != 0:
                    raise Exception(f"分块上传失败(slice {slice_id}): {r.get('message', '未知错误')}")
            finally:
                os.remove(tmp_path)
            slice_id += 1

    # 3) 完成分块上传
    body = {"request_id": str(uuid.uuid4()), "app_id": app_id, "upload_id": upload_id}
    body_bytes = json.dumps(body).encode("utf-8")
    headers = ost_build_auth(OST_MPCOMPLETE_URL, "POST", api_key, api_secret, body_bytes)
    headers["content-type"] = "application/json"
    resp = requests.post(OST_MPCOMPLETE_URL, headers=headers, data=body_bytes, timeout=30)
    result = resp.json()
    print(f"[DEBUG] OST mpcomplete response: {json.dumps(result, ensure_ascii=False)}")
    if result.get("code") != 0:
        raise Exception(f"分块上传完成失败: {result.get('message', '未知错误')}")

    return result["data"]["url"]


def ost_create_task(app_id, api_key, api_secret, audio_url, audio_format, duration_ms=0):
    """极速版：创建转写任务"""
    if audio_format in ("wav", "pcm"):
        encoding = "raw"
    elif audio_format == "mp3":
        encoding = "lame"
    else:
        encoding = "raw"

    request_id = str(uuid.uuid4())
    body = {
        "common": {"app_id": app_id},
        "business": {
            "request_id": request_id,
            "language": "zh_cn",
            "domain": "pro_ost_ed",
            "accent": "mandarin",
            "vspp_on": 1,
            "speaker_num": 0,
            "duration": duration_ms // 1000 if duration_ms else 0,
        },
        "data": {
            "audio_url": audio_url,
            "audio_src": "http",
            "format": "16k",
            "encoding": encoding,
        },
    }

    body_bytes = json.dumps(body).encode("utf-8")
    headers = ost_build_auth(OST_CREATE_URL, "POST", api_key, api_secret, body_bytes)
    headers["content-type"] = "application/json"

    resp = requests.post(OST_CREATE_URL, headers=headers, data=body_bytes)
    result = resp.json()
    print(f"[DEBUG] OST create task response: {json.dumps(result, ensure_ascii=False)}")

    if result.get("code") != 0:
        raise Exception(f"创建任务失败: {result.get('message', '未知错误')}")

    return result["data"]["task_id"]


def ost_query_task(app_id, api_key, api_secret, ost_task_id):
    """极速版：查询任务状态，返回 (status, error_msg, result_data)"""
    body = {
        "common": {"app_id": app_id},
        "business": {"task_id": ost_task_id},
    }

    body_bytes = json.dumps(body).encode("utf-8")
    headers = ost_build_auth(OST_QUERY_URL, "POST", api_key, api_secret, body_bytes)
    headers["content-type"] = "application/json"

    resp = requests.post(OST_QUERY_URL, headers=headers, data=body_bytes, timeout=30)
    result = resp.json()

    if result.get("code") != 0:
        return "failed", result.get("message", "查询失败"), None

    data = result.get("data", {})
    task_status = str(data.get("task_status", "1"))

    if task_status in ("3", "4"):
        return "done", None, data.get("result", {})
    elif task_status == "2":
        return "processing", None, None
    else:
        return "pending", None, None


# ========== 后台轮询线程 ==========

def poll_task(task_id, api_key, api_secret, order_id):
    """后台轮询转写结果"""
    start = time.time()
    while True:
        try:
            status, content = poll_result(api_key, api_secret, order_id)
            elapsed = int(time.time() - start)

            with tasks_lock:
                if task_id not in tasks:
                    break
                tasks[task_id]["elapsed"] = elapsed

                if status == 4:
                    segments = parse_result(content)
                    tasks[task_id]["status"] = "done"
                    tasks[task_id]["segments"] = segments
                    # 保存原始 orderResult 用于调试
                    order_result_raw = content.get("orderResult", "")
                    if order_result_raw:
                        debug_path = os.path.join(RESULT_DIR, f"{task_id}_raw.json")
                        with open(debug_path, "w", encoding="utf-8") as df:
                            if isinstance(order_result_raw, str):
                                df.write(order_result_raw)
                            else:
                                json.dump(order_result_raw, df, ensure_ascii=False, indent=2)
                    # 保存结果到文件
                    save_result(task_id, tasks[task_id])
                    break
                elif status == -1:
                    fail_type = content.get("orderInfo", {}).get("failType", "未知")
                    tasks[task_id]["status"] = "failed"
                    tasks[task_id]["error"] = f"转写失败 (failType: {fail_type})"
                    break
                elif status == 0:
                    tasks[task_id]["status_text"] = "已创建，等待处理"
                elif status == 3:
                    tasks[task_id]["status_text"] = "处理中..."

        except Exception as e:
            elapsed = int(time.time() - start)
            with tasks_lock:
                if task_id in tasks:
                    tasks[task_id]["elapsed"] = elapsed
            # SSL 等临时错误，继续重试
            pass

        time.sleep(15)


def poll_task_ost(task_id, app_id, api_key, api_secret, ost_task_id):
    """极速版：后台轮询转写结果"""
    start = time.time()
    while True:
        try:
            status, error_msg, result_data = ost_query_task(
                app_id, api_key, api_secret, ost_task_id
            )
            elapsed = int(time.time() - start)

            with tasks_lock:
                if task_id not in tasks:
                    break
                tasks[task_id]["elapsed"] = elapsed

                if status == "done":
                    lattice = result_data.get("lattice", [])
                    segments = parse_result(
                        {"orderResult": json.dumps({"lattice": lattice})}
                    )
                    tasks[task_id]["status"] = "done"
                    tasks[task_id]["segments"] = segments
                    raw_path = os.path.join(RESULT_DIR, f"{task_id}_raw.json")
                    with open(raw_path, "w", encoding="utf-8") as df:
                        json.dump(result_data, df, ensure_ascii=False, indent=2)
                    save_result(task_id, tasks[task_id])
                    break
                elif status == "failed":
                    tasks[task_id]["status"] = "failed"
                    tasks[task_id]["error"] = error_msg or "转写失败"
                    break
                elif status == "processing":
                    tasks[task_id]["status_text"] = "处理中..."
                else:
                    tasks[task_id]["status_text"] = "已创建，等待处理"

        except Exception as e:
            elapsed = int(time.time() - start)
            with tasks_lock:
                if task_id in tasks:
                    tasks[task_id]["elapsed"] = elapsed

        time.sleep(5)


def save_result(task_id, task_data):
    """保存结果到文件"""
    filepath = os.path.join(RESULT_DIR, f"{task_id}.json")
    data = {
        "task_id": task_id,
        "filename": task_data.get("filename", ""),
        "created_at": task_data.get("created_at", ""),
        "segments": task_data.get("segments", []),
    }
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_result(task_id):
    """从文件加载结果"""
    filepath = os.path.join(RESULT_DIR, f"{task_id}.json")
    if not os.path.exists(filepath):
        return None
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def load_all_results():
    """加载所有历史结果"""
    results = []
    if not os.path.exists(RESULT_DIR):
        return results
    for fname in sorted(os.listdir(RESULT_DIR), reverse=True):
        if fname.endswith(".json"):
            filepath = os.path.join(RESULT_DIR, fname)
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                results.append(data)
    return results


# ========== 导出功能 ==========

def export_txt(task_data, speaker_names):
    """导出为 TXT"""
    filename = task_data.get("filename", "转写结果")
    base_name = os.path.splitext(filename)[0]
    lines = [f"{base_name} 录音转写结果", "=" * 40, ""]
    for rl, text in task_data.get("segments", []):
        name = speaker_names.get(str(rl), f"说话人{rl}")
        lines.append(f"【{name}】{text}")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_docx(task_data, speaker_names):
    """导出为 Word"""
    doc = Document()
    filename = task_data.get("filename", "转写结果")
    base_name = os.path.splitext(filename)[0]

    title = doc.add_heading(base_name + " 录音转写结果", level=1)
    for run in title.runs:
        run.font.size = Pt(18)

    for rl, text in task_data.get("segments", []):
        name = speaker_names.get(str(rl), f"说话人{rl}")
        p = doc.add_paragraph()
        run = p.add_run(f"【{name}】")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_pdf(task_data, speaker_names):
    """导出为 PDF（已移除）"""
    raise Exception("PDF 导出功能已移除")


# ========== 认证装饰器 ==========

def login_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


# ========== 路由 ==========

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        password = request.form.get("password", "")
        if password == APP_PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("index"))
        return render_template("login.html", error="密码错误")
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/")
@login_required
def index():
    return render_template("index.html")


@app.route("/api/upload", methods=["POST"])
@login_required
def api_upload():
    """上传音频并启动转写"""
    provider = request.form.get("provider", "xfyun")
    app_id = request.form.get("app_id", "").strip()
    api_key = request.form.get("api_key", "").strip()
    api_secret = request.form.get("api_secret", "").strip()

    if not all([app_id, api_key, api_secret]):
        return jsonify({"error": "请填写完整的讯飞认证信息"}), 400

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "请选择音频文件"}), 400

    # 保存文件
    task_id = str(uuid.uuid4())[:8]
    ext = os.path.splitext(file.filename)[1]
    filepath = os.path.join(UPLOAD_DIR, f"{task_id}{ext}")
    file.save(filepath)

    try:
        if provider == "xfyun_speed":
            duration_ms = get_audio_duration_ms(filepath)
            audio_url, audio_fmt = ost_upload_file(
                app_id, api_key, api_secret, filepath
            )
            ost_task_id = ost_create_task(
                app_id, api_key, api_secret, audio_url, audio_fmt, duration_ms
            )
            estimate_time = max(20000, duration_ms * 20)
        else:
            order_id, estimate_time = upload_audio(
                app_id, api_key, api_secret, filepath
            )
    except Exception as e:
        os.remove(filepath)
        return jsonify({"error": str(e)}), 500

    with tasks_lock:
        tasks[task_id] = {
            "task_id": task_id,
            "order_id": ost_task_id if provider == "xfyun_speed" else order_id,
            "filename": file.filename,
            "status": "processing",
            "status_text": "已上传，等待处理",
            "estimate_time": estimate_time,
            "elapsed": 0,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "provider": provider,
        }

    if provider == "xfyun_speed":
        t = threading.Thread(
            target=poll_task_ost,
            args=(task_id, app_id, api_key, api_secret, ost_task_id),
            daemon=True,
        )
    else:
        t = threading.Thread(
            target=poll_task,
            args=(task_id, api_key, api_secret, order_id),
            daemon=True,
        )
    t.start()

    return jsonify({"task_id": task_id, "estimate_time": estimate_time})


@app.route("/api/status/<task_id>")
@login_required
def api_status(task_id):
    """查询转写状态"""
    with tasks_lock:
        task = tasks.get(task_id)
    if not task:
        # 尝试从文件加载
        data = load_result(task_id)
        if data:
            return jsonify({
                "status": "done",
                "task_id": task_id,
                "filename": data.get("filename", ""),
                "segment_count": len(data.get("segments", [])),
            })
        return jsonify({"error": "任务不存在"}), 404

    return jsonify({
        "status": task["status"],
        "task_id": task_id,
        "status_text": task.get("status_text", ""),
        "elapsed": task.get("elapsed", 0),
        "estimate_time": task.get("estimate_time", 0),
        "error": task.get("error", ""),
    })


@app.route("/result/<task_id>")
@login_required
def result_page(task_id):
    """结果展示页"""
    data = load_result(task_id)
    if not data:
        return redirect(url_for("index"))

    # 提取说话人列表
    speakers = sorted(set(str(rl) for rl, _ in data.get("segments", [])))
    single_speaker = len(speakers) <= 1 and len(data.get("segments", [])) > 0
    return render_template("result.html", task=data, speakers=speakers, task_id=task_id, single_speaker=single_speaker)


@app.route("/api/result/<task_id>")
@login_required
def api_result(task_id):
    """获取转写结果 JSON"""
    data = load_result(task_id)
    if not data:
        return jsonify({"error": "结果不存在"}), 404
    return jsonify(data)


@app.route("/api/export/<task_id>/<format>")
@login_required
def api_export(task_id, format):
    """导出结果"""
    data = load_result(task_id)
    if not data:
        return jsonify({"error": "结果不存在"}), 404

    # 获取说话人名称映射
    speaker_names = {}
    for key, value in request.args.items():
        if key.startswith("speaker_"):
            speaker_names[key.replace("speaker_", "")] = value

    base_name = os.path.splitext(data.get("filename", "转写结果"))[0]

    if format == "txt":
        content = export_txt(data, speaker_names)
        import io
        return send_file(
            io.BytesIO(content),
            mimetype="text/plain",
            as_attachment=True,
            download_name=f"{base_name}_转写结果.txt",
        )
    elif format == "docx":
        buf = export_docx(data, speaker_names)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{base_name}_转写结果.docx",
        )
    else:
        return jsonify({"error": "不支持的格式"}), 400


@app.route("/history")
@login_required
def history_page():
    """历史记录页"""
    return render_template("history.html")


if __name__ == "__main__":
    print(f"录音转写 Web 应用启动")
    print(f"访问地址: http://localhost:{PORT}")
    print(f"登录密码: {APP_PASSWORD}")
    app.run(host="0.0.0.0", port=PORT, debug=False)
